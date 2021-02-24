import hashlib
import datetime
import re
import time
import calendar
import xlrd
from dateutil.relativedelta import relativedelta
from datetime import timedelta
from docx import Document
# from pandas.core.dtypes.inference import is_number
from collections import Counter
import os
from django.shortcuts import render
import requests
import pymysql
import json
from django.http import HttpResponse, JsonResponse
# from .models import form_ShiPin_Report
# from datetime import datetime, timedelta
from django.db.models import Q
import pymysql
from .models import video_user as User
from .models import FormShipinReport as ShiPin
from .models import *
# Create your views here.


def index(request):
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from form_ShiPin_Report ORDER BY ID DESC limit 8"
    cursor.execute(name_sql)
    result = cursor.fetchall()
    db.commit()
    db.close()

    data = {'data': []}
    nvdan = []
    nandan = []
    nvshuang = []
    nanshuang = []
    hunshuang = []
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from player_list"
    cursor.execute(name_sql)
    results = cursor.fetchall()
    db.commit()
    db.close()
    for i in results:
        j = list(i)
        if j[2] == '女单':
            nvdan.append(j)
        elif j[2] == '男单':
            nandan.append(j)
        elif j[2] == '女双':
            nvshuang.append(j)
        elif j[2] == '男双':
            nanshuang.append(j)
        elif j[2] == '混双':
            hunshuang.append(j)
    data['player1'] = nvdan[:3] + nandan[:3]
    data['player2'] = nvshuang[:2] + nanshuang[:2] + hunshuang[:2]

    matchs = []
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from match_list order by match_date DESC"
    cursor.execute(name_sql)
    results = cursor.fetchall()
    db.commit()
    db.close()
    for i in results:
        j = list(i)
        matchs.append(j)

    match1 = matchs[:6]
    match2 = matchs[6:12]
    data['match1'] = match1
    data['match2'] = match2

    for i in result:
        d1 = {
            'name': '-'.join(i[7].split('-')[3:]),
            'url': i[3]
        }
        data['data'].append(d1)
    response = HttpResponse(json.dumps(data))
    # response["Access-Control-Allow-Origin"] = "*"
    # response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    # response["Access-Control-Max-Age"] = "1000"
    # response["Access-Control-Allow-Headers"] = " * "
    return response


def player_video(request):
    data = {'data': []}
    player = request.GET.get('player')
    match = ShiPin.objects.filter(Q(single_line_shipinmingcheng__icontains=player))
    for i in match:
        dd = {
            'name': '-'.join(i.single_line_shipinmingcheng.split('-')[3:]),
            'url': i.url
        }
        data['data'].append(dd)
    response = HttpResponse(json.dumps(data))
    # response["Access-Control-Allow-Origin"] = "*"
    # response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    # response["Access-Control-Max-Age"] = "1000"
    # response["Access-Control-Allow-Headers"] = " * "
    return response


def project_video(request):
    data = {'data': []}
    match = request.GET.get('match')
    year = request.GET.get('year')
    project = request.GET.get('project')
    results = ShiPin.objects.filter(Q(single_line_shipinmingcheng__icontains=match))
    for i in results:
        n = i.single_line_shipinmingcheng.split('-')
        if n[0] == year:
            if int(n[5]) != 9 and n[4] == project:
                dd = {
                    'name': '-'.join(i.single_line_shipinmingcheng.split('-')[3:]),
                    'url': i.url
                }
                data['data'].append(dd)
    response = HttpResponse(json.dumps(data))
    # response["Access-Control-Allow-Origin"] = "*"
    # response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    # response["Access-Control-Max-Age"] = "1000"
    # response["Access-Control-Allow-Headers"] = " * "
    return response


def playercn_info(request):
    data = {}
    player = request.GET.get('player')
    xinxi = [player, 0, 0, 0, 0, 0, 0, 0, 0]
    playerid = FormYundongyuanReport.objects.filter(single_line_zhongwenming=player)[0].id
    player_yingwenming = FormYundongyuanReport.objects.filter(single_line_zhongwenming=player)[0].single_line_yingwenming
    res_sjpm = FormPaiming.objects.filter(Q(form_yundongyuan_id=playerid) & Q(dropdown_paimingfenlei='世界排名'))
    sjpms = []
    for i in res_sjpm:
        pmri = int(i.formula_riqi)
        pm = i.number_paiming
        sjpms.append([pmri, pm])
    sjpms.sort(key=lambda x:x[0], reverse=True)
    sjpm = sjpms[0][1]

    res_zjspm = FormPaiming.objects.filter(Q(form_yundongyuan_id=playerid) & Q(dropdown_paimingfenlei='总决赛排名'))
    zjspms = []
    for i in res_zjspm:
        pmri = int(i.formula_riqi)
        pm = i.number_paiming
        zjspms.append([pmri, pm])
    zjspms.sort(key=lambda x: x[0], reverse=True)
    zjspm = zjspms[0][1]

    res_jf_2020 = FormJifen.objects.filter(Q(form_yundongyuan=player_yingwenming) & Q(number_nian='2020'))
    mc_2020_1 = 0
    mc_2020_2 = 0
    mc_2020_3 = 0
    for i in res_jf_2020:
        if i.dropdown_mingci == '1':
            mc_2020_1 += 1
        elif i.dropdown_mingci == '2':
            mc_2020_2 += 1
        elif i.dropdown_mingci == '3/4':
            mc_2020_3 += 1

    res_jf_2021 = FormJifen.objects.filter(Q(form_yundongyuan=player_yingwenming) & Q(number_nian='2021'))
    mc_2021_1 = 0
    mc_2021_2 = 0
    mc_2021_3 = 0
    for i in res_jf_2021:
        if i.dropdown_mingci == '1':
            mc_2021_1 += 1
        elif i.dropdown_mingci == '2':
            mc_2021_2 += 1
        elif i.dropdown_mingci == '3/4':
            mc_2021_3 += 1
    xinxi[1] = sjpm
    xinxi[2] = zjspm
    xinxi[3] = mc_2021_1
    xinxi[4] = mc_2021_2
    xinxi[5] = mc_2021_3
    xinxi[6] = mc_2020_1
    xinxi[7] = mc_2020_2
    xinxi[8] = mc_2020_3

    jctn = [Form3000M, Form30M, FormBmi, FormWotuiMax, FormZuoweitiqianqu, FormChuizhizongtiao, FormYintixiangshang, FormShendunMax, FormBeijinaili, FormFujinaili]
    zxtn = [Form400Mx5, FormLidingtiaoyuan, FormYuandierjitiaoZuo, FormYuandierjitiaoYou, Form5000M, FormShendun90Kg, FormShuangyao1500, FormSijiaopao20]
    pinfen_jctn = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],
                   [0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]
    pinfen_zxtn = [[0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0],
                   [0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0]]
    pinfen_fms = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0],
                  [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]
    times_jctn = []
    data_jctn = []
    for i in jctn:
        result = i.objects.filter(form_yundongyuan=player)
        for j in result:
            # print(j)
            try:
                t1_cs = j.date_field_ceshi.split('-')
            except Exception as e:
                t1_cs = j.date_field_ceshiriqi.split('-')
            y = str(t1_cs[2])
            m = str(list(calendar.month_abbr).index(t1_cs[1])).zfill(2)
            d = str(t1_cs[0])
            t1 = int(y + m + d)
            if t1 not in times_jctn:
                times_jctn.append(t1)
            try:
                d1 = [t1, str(i).split("'")[1].split('.')[2], j.dropdown_pingfen]
            except Exception as e:
                d1 = [t1, str(i).split("'")[1].split('.')[2], j.number_pingfen]
            if d1 not in data_jctn:
                data_jctn.append(d1)
    times_jctn.sort(reverse=True)
    times_jctn = times_jctn[:6]
    for t_jctn in times_jctn:
        n = times_jctn.index(t_jctn)
        for j in data_jctn:
            if j[2] != '':
                pf_jctn = int(j[2])
            else:
                pf_jctn = 0
            if t_jctn == j[0]:
                if j[1] == 'Form3000M':
                    pinfen_jctn[n][0] = pf_jctn
                elif j[1] == 'Form30M':
                    pinfen_jctn[n][1] = pf_jctn
                elif j[1] == 'FormBmi':
                    pinfen_jctn[n][2] = pf_jctn
                elif j[1] == 'FormWotuiMax':
                    pinfen_jctn[n][3] = pf_jctn
                elif j[1] == 'FormZuoweitiqianqu':
                    pinfen_jctn[n][4] = pf_jctn
                elif j[1] == 'FormChuizhizongtiao':
                    pinfen_jctn[n][5] = pf_jctn
                elif j[1] == 'FormYintixiangshang':
                    pinfen_jctn[n][6] = pf_jctn
                elif j[1] == 'FormShendunMax':
                    pinfen_jctn[n][7] = pf_jctn
                elif j[1] == 'FormBeijinaili':
                    pinfen_jctn[n][8] = pf_jctn
                elif j[1] == 'FormFujinaili':
                    pinfen_jctn[n][9] = pf_jctn
    times_zxtn = []
    data_zxtn = []
    for i in zxtn:
        result = i.objects.filter(form_yundongyuan=player)
        pf = []
        for j in result:
            # print(j)
            try:
                t1_cs = j.date_field_ceshi.split('-')
            except Exception as e:
                t1_cs = j.date_field_ceshiriqi.split('-')
            y = str(t1_cs[2])
            m = str(list(calendar.month_abbr).index(t1_cs[1])).zfill(2)
            d = str(t1_cs[0])
            t1 = int(y + m + d)
            if t1 not in times_zxtn:
                times_zxtn.append(t1)
            try:
                d1 = [t1, str(i).split("'")[1].split('.')[2], j.dropdown_pingfen]
            except Exception as e:
                d1 = [t1, str(i).split("'")[1].split('.')[2], j.number_pingfen]
            if d1 not in data_zxtn:
                data_zxtn.append(d1)
    times_zxtn.sort(reverse=True)
    times_zxtn = times_zxtn[:6]
    for t_zxtn in times_zxtn:
        n = times_zxtn.index(t_zxtn)
        for j in data_zxtn:
            if j[2] != '':
                pf_zxtn = int(j[2])
            else:
                pf_zxtn = 0
            if t_zxtn == j[0]:
                if j[1] == 'Form400Mx5':
                    pinfen_zxtn[n][0] = pf_zxtn
                elif j[1] == 'FormLidingtiaoyuan':
                    pinfen_zxtn[n][1] = pf_zxtn
                elif j[1] == 'FormYuandierjitiaoZuo':
                    pinfen_zxtn[n][2] = pf_zxtn
                elif j[1] == 'FormYuandierjitiaoYou':
                    pinfen_zxtn[n][3] = pf_zxtn
                elif j[1] == 'Form5000M':
                    pinfen_zxtn[n][4] = pf_zxtn
                elif j[1] == 'FormShendun90Kg':
                    pinfen_zxtn[n][5] = pf_zxtn
                elif j[1] == 'FormShuangyao1500':
                    pinfen_zxtn[n][6] = pf_zxtn
                elif j[1] == 'FormSijiaopao20':
                    pinfen_zxtn[n][7] = pf_zxtn


    times_fms = []
    data_fms = []
    fms = ['上踏步（左）', '上踏步（右）', '直线弓箭步（左）', '直线弓箭步（右）', '肩部灵活性（左）', '肩部灵活性（右）', '直腿主动上抬（左）', '直腿主动上抬（右）', '深蹲',
           '躯干稳定俯卧撑', '扭转稳定性（左）', '扭转稳定性（右）']
    for i in fms:
        result = Fms.objects.filter(Q(form_yundongyuan=player) & Q(dropdown_dongzuo=i))
        for j in result:
            try:
                t1_cs = j.date_field_ceshi.split('-')
            except Exception as e:
                t1_cs = j.date_field_ceshiriqi.split('-')
            y = str(t1_cs[2])
            m = str(list(calendar.month_abbr).index(t1_cs[1])).zfill(2)
            d = str(t1_cs[0])
            t1 = int(y + m + d)
            if t1 not in times_fms:
                times_fms.append(t1)
            d1 = [t1, i, j.number_pingfen]
            if d1 not in data_fms:
                data_fms.append(d1)
    times_fms.sort(reverse=True)
    times_fms = times_fms[:6]
    for t_fms in times_fms:
        n = times_fms.index(t_fms)
        for j in data_fms:
            if j[2] != '':
                pf_fms = int(j[2])
            else:
                pf_fms = 0
            if t_fms == j[0]:
                if j[1] == '上踏步（左）':
                    pinfen_fms[n][0] = pf_fms
                elif j[1] == '上踏步（右）':
                    pinfen_fms[n][1] = pf_fms
                elif j[1] == '直线弓箭步（左）':
                    pinfen_fms[n][2] = pf_fms
                elif j[1] == '直线弓箭步（右）':
                    pinfen_fms[n][3] = pf_fms
                elif j[1] == '肩部灵活性（左）':
                    pinfen_fms[n][4] = pf_fms
                elif j[1] == '肩部灵活性（右）':
                    pinfen_fms[n][5] = pf_fms
                elif j[1] == '直腿主动上抬（左）':
                    pinfen_fms[n][6] = pf_fms
                elif j[1] == '直腿主动上抬（右）':
                    pinfen_fms[n][7] = pf_fms
                elif j[1] == '深蹲':
                    pinfen_fms[n][8] = pf_fms
                elif j[1] == '躯干稳定俯卧撑':
                    pinfen_fms[n][9] = pf_fms
                elif j[1] == '扭转稳定性（左）':
                    pinfen_fms[n][10] = pf_fms
                elif j[1] == '扭转稳定性（右）':
                    pinfen_fms[n][11] = pf_fms

    res_xxl = FormXunlianliang.objects.filter(form_yundongyuan=player)
    result_xxl = []
    keys = []
    for i in res_xxl:
        xx_time = i.date_field_xunlianriqi.split('-')
        y = str(xx_time[2])
        m = str(list(calendar.month_abbr).index(xx_time[1])).zfill(2)
        d = str(xx_time[0])
        t1 = y+'-'+m+'-'+d
        t2 = int(y+m+d)
        result_xxl.append([t2, i.dropdown_xunliankemu, int(i.number_xunlianshichang)])
        if t1 not in keys:
            keys.append(t1)
    keys.sort(reverse=True)
    max_time = keys[0]
    # max_time2 = keys[7]
    # max_time3 = keys[14]
    # max_time4 = keys[21]
    # max_time5 = keys[28]
    # max_time6 = keys[35]
    # print(keys)
    # print(max_time1, max_time2, max_time3, max_time4, max_time5, max_time6)
    key_list = []
    ind = 1
    keys.append('2200-01-01')
    while ind < len(keys)-1:
        # print(int(time.mktime(time.strptime(keys[ind], '%Y-%m-%d')))-int(time.mktime(time.strptime(keys[ind-1], '%Y-%m-%d'))))
        # print(int(time.mktime(time.strptime(keys[ind-1], '%Y-%m-%d'))))
        if int(time.mktime(time.strptime(keys[ind], '%Y-%m-%d'))) - int(time.mktime(time.strptime(keys[ind-1], '%Y-%m-%d'))) == -86400:
            start = ind - 1
            while int(time.mktime(time.strptime(keys[ind], '%Y-%m-%d'))) - int(time.mktime(time.strptime(keys[ind-1], '%Y-%m-%d'))) == -86400:
                ind += 1
            k1 = int(''.join(keys[start].split('-')))
            key_list.append(k1)
        else:
            ind += 1
    # for i in keys:
    #     format_time = int(time.mktime(time.strptime(i, '%Y-%m-%d')))
    shichang_xxl = []
    for i in key_list:
        shichang = [0, 0, 0, 0]
        jzs = []
        xjs = []
        tnzx = []
        tnll = []
        for j in result_xxl:
            if i-6 <= j[0] <= i:
                if j[1] == '技术（技战术）':
                    jzs.append(j[2])
                    # shichang[0] = j[2]
                elif j[1] == '技术（小技术）':
                    xjs.append(j[2])
                    # shichang[1] = j[2]
                elif j[1] == '体能（专项）':
                    tnzx.append(j[2])
                    # shichang[2] = j[2]
                elif j[1] == '体能（力量）':
                    tnll.append(j[2])
                    # shichang[3] = j[2]
        shichang[0] = sum(jzs)
        shichang[1] = sum(xjs)
        shichang[2] = sum(tnzx)
        shichang[3] = sum(tnll)
        shichang_xxl.append(shichang)
    while len(shichang_xxl) < 6:
        shichang_xxl.append([0, 0, 0, 0])
    res_slsh = FormShenglishenghua.objects.filter(form_yundongyuan=player)
    result_slsh = []
    t_list = []
    for i in res_slsh:
        xx_time = i.date_field_ceshiriqi.split('-')
        y = str(xx_time[2])
        m = str(list(calendar.month_abbr).index(xx_time[1])).zfill(2)
        d = str(xx_time[0])
        t1 = int(y + m + d)
        result_slsh.append([t1, i.formula_ceshixiangmu, float(i.decimal_ceshijieguo)])
        if t1 not in t_list:
            t_list.append(t1)
    t_list.sort(reverse=True)
    t_list = t_list[:6]
    slsh_t = []
    for i in t_list:
        ts1 = str(i)
        ts2 = ts1[:4] + '-' + ts1[4:6] + '-' + ts1[6:8]
        slsh_t.append(ts2)
    r1_slsh = []
    for i in result_slsh:
        if i[0] in t_list:
            r1_slsh.append(i)
    # r1_slsh.sort(key=lambda x:x[0], reverse=True)
    csjg_slsh = [[0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0], [0, 0, 0, 0, 0, 0]]
    bxb = []
    pzc = []
    gt = []
    hxb = []
    jsjm = []
    xns = []
    xqyj = []
    xhdb = []
    for i in r1_slsh:
        if i[1] == '白细胞':
            bxb.append(i)
        elif i[1] == '皮质醇':
            pzc.append(i)
        elif i[1] == '睾酮':
            gt.append(i)
        elif i[1] == '红细胞':
            hxb.append(i)
        elif i[1] == '肌酸激酶':
            jsjm.append(i)
        elif i[1] == '血尿素':
            xns.append(i)
        elif i[1] == '血球压积':
            xqyj.append(i)
        elif i[1] == '血红蛋白':
            xhdb.append(i)
    slsh_all = [bxb, pzc, gt, hxb, jsjm, xns, xqyj, xhdb]
    for o in slsh_all:
        m = slsh_all.index(o)
        for i in o:
            for t in t_list:
                n = t_list.index(t)
                if i[0] == t:
                    csjg_slsh[m][n] = i[2]
    jc_times = []
    zx_times = []
    fms_times = []
    xxl_times = []
    for jc in times_jctn:
        tt1 = str(jc)
        tt2 = tt1[:4] + '-' + tt1[4:6] + '-' + tt1[6:8]
        jc_times.append(tt2)
    for zx in times_zxtn:
        tt1 = str(zx)
        tt2 = tt1[:4] + '-' + tt1[4:6] + '-' + tt1[6:8]
        zx_times.append(tt2)
    for ft in times_fms:
        tt1 = str(ft)
        tt2 = tt1[:4] + '-' + tt1[4:6] + '-' + tt1[6:8]
        fms_times.append(tt2)
    for xt in key_list:
        tt1 = str(xt)
        tt2 = datetime.date(int(tt1[:4]), int(tt1[4:6]), int(tt1[6:8])).isocalendar()
        tt3 = str(tt2[0]) + '年' + str(tt2[1]) + '周'
        xxl_times.append(tt3)
    while len(jc_times) < 6:
        jc_times.append(['-'])
    while len(zx_times) < 6:
        zx_times.append(['-'])
    while len(fms_times) < 6:
        fms_times.append(['-'])
    while len(xxl_times) < 6:
        xxl_times.append(['-'])
    while len(slsh_t) < 6:
        slsh_t.append(['-'])

    res_shenglvdanda = ShenglvDanda.objects.filter(form_yundongyuan_single_line_zhongwenming=player)
    res_shenglvzuhe = ShenglvZuhe.objects.filter(form_shuangdazuhe_zhongguo_single_line1__contains=player)
    wgqy = []
    shenglv = []
    for i in res_shenglvdanda:
        ds = i.form_yundongyuan_8_single_line_zhongwenming
        sl = i.percent_shenglv
        wgqy.append({'name': ds, 'max': 100})
        shenglv.append(sl)
    for i in res_shenglvzuhe:
        ds = i.form_shuangdazuhe_8_single_line1
        sl = i.percent_shenglv
        wgqy.append({'name': ds, 'max': 100})
        shenglv.append(sl)

    if len(wgqy) == 0:
        wgqy.append({'name': '-', 'max': 100})
        shenglv.append('0')

    res_shipin = FormShipinReport.objects.filter(single_line_shipinmingcheng__contains=player)
    shipin_list = []
    shipin_lista = []
    for i in res_shipin:
        spmc = i.single_line_shipinmingcheng
        ss = i.query_saishimingcheng
        lunci = i.dropdown_lunci
        zbf = i.dropdown_bisaijieguo
        mjbf = i.single_line_bifen
        shipin = [spmc, ss, lunci, zbf, mjbf]
        shipin_lista.append(shipin)
    shipin_lista.sort(key=lambda x: x[0], reverse=True)
    for i in shipin_lista:
        spmc = '-'.join(i[0].split('-')[6:])
        ss = i[1]
        lunci = i[2]
        zbf = i[3]
        mjbf = i[4]
        shipin = [spmc, ss, lunci, zbf, mjbf]
        shipin_list.append(shipin)
    # shipin.sort(keys=lambda x: x[])
    data['xinxi'] = xinxi
    data['jctn_time'] = jc_times
    data['jctn'] = pinfen_jctn
    data['zxtn_time'] = zx_times
    data['zxtn'] = pinfen_zxtn
    data['fms_time'] = fms_times
    data['fms'] = pinfen_fms
    data['xxl_time'] = xxl_times
    data['xxl'] = shichang_xxl
    data['slsh_time'] = slsh_t
    data['slsh'] = csjg_slsh
    data['wgqy'] = wgqy
    data['shenglv'] = shenglv
    data['shipin'] = shipin_list
    # response = HttpResponse(json.dumps(data))
    response = HttpResponse(json.dumps(data))
    return response


def search(request):
    data = {'data': []}
    content = request.GET.get('content').split(" ")
    ln = []
    listcon = []
    list1 = []
    list2 = []
    for i in content:
        # print(i)
        if i.encode('UTF-8').isalpha() == True:
            ln.append(i)
        else:
            listcon.append(i)
    a = ''.join(ln)
    if len(a) != 0:
        db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
        cursor = db.cursor()
        name_sql = "select Single_Line_ZhongWenMing from form_YunDongYuan_Report where Single_Line_YingWenMing='" + a + "'"
        cursor.execute(name_sql)
        result = cursor.fetchone()
        db.commit()
        db.close()
        for i in result:
            listcon.append(i)
    for i in listcon:
        # print(i)
        match = ShiPin.objects.filter(Q(single_line_shipinmingcheng__icontains=i))
        for h in match:
            list1.append(h)
    if len(listcon) == 1:
        for i in list1:
            if i not in list2:
                list2.append(i)
    elif len(listcon) == 2:
        l1 = dict(Counter(list1))
        list2 = [key for key, value in l1.items() if 3 > value > 1]
        # print(l1)
        # print(list2)
    elif len(listcon) == 3:
        l1 = dict(Counter(list1))
        list2 = [key for key, value in l1.items() if 4 > value > 2]
    elif len(listcon) == 4:
        l1 = dict(Counter(list1))
        list2 = [key for key, value in l1.items() if 5 > value > 3]
    elif len(listcon) == 5:
        l1 = dict(Counter(list1))
        list2 = [key for key, value in l1.items() if 6 > value > 4]
    for i in list2:
        # n = i.single_line_shipinmingcheng.split('-')
        dd = {
            'name': '-'.join(i.single_line_shipinmingcheng.split('-')[3:]),
            'url': i.url
        }
        data['data'].append(dd)

    response = HttpResponse(json.dumps(data))
    # response["Access-Control-Allow-Origin"] = "*"
    # response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    # response["Access-Control-Max-Age"] = "1000"
    # response["Access-Control-Allow-Headers"] = " * "
    return response


def login(request):
    if request.method == 'POST':
        name = request.POST.get('user')
        pwd = request.POST.get('pwd')
        result = User.objects.filter(username=name)[0]
        if result.password == pwd:
            data = {'data': name}
        else:
            data = {'data': ''}
        # if name == 'my':
        # request.session['name'] = name
        # response = HttpResponse('set cookie')

        # data = {'data': name}
        response = HttpResponse(json.dumps(data))

        # response.set_cookie("name", 'my')
        # response["Access-Control-Allow-Origin"] = "null"
        # response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
        # response["Access-Control-Max-Age"] = "1000"
        # response["Access-Control-Allow-Headers"] = " * "
        return response


def player_list(request):
    data = {}
    nvdan = []
    nandan = []
    nvshuang = []
    nanshuang = []
    hunshuang = []
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from player_list"
    cursor.execute(name_sql)
    results = cursor.fetchall()
    db.commit()
    db.close()
    for i in results:
        j = list(i)
        if j[2] == '女单':
            nvdan.append(j)
        elif j[2] == '男单':
            nandan.append(j)
        elif j[2] == '女双':
            nvshuang.append(j)
        elif j[2] == '男双':
            nanshuang.append(j)
        elif j[2] == '混双':
            hunshuang.append(j)
    data['nvdan'] = nvdan
    data['nandan'] = nandan
    data['nvshuang'] = nvshuang
    data['nanshuang'] = nanshuang
    data['hunshuang'] = hunshuang
    response = HttpResponse(json.dumps(data))
    return response


def playercn_list(request):
    data = {}
    nvdan = []
    nandan = []
    nvshuang = []
    nanshuang = []
    hunshuang = []
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from playercn_list"
    cursor.execute(name_sql)
    results = cursor.fetchall()
    db.commit()
    db.close()
    for i in results:
        j = list(i)
        if j[2] == '女单':
            nvdan.append(j)
        elif j[2] == '男单':
            nandan.append(j)
        elif j[2] == '女双':
            nvshuang.append(j)
        elif j[2] == '男双':
            nanshuang.append(j)
        elif j[2] == '混双':
            hunshuang.append(j)
    data['nvdan'] = nvdan
    data['nandan'] = nandan
    data['nvshuang'] = nvshuang
    data['nanshuang'] = nanshuang
    data['hunshuang'] = hunshuang
    response = HttpResponse(json.dumps(data))
    return response


def match_list(request):
    data = {}
    matchs = []
    db = pymysql.connect('video.hbang.com.cn', 'video', 'P@ssw0rd235', 'video')
    cursor = db.cursor()
    name_sql = "select * from match_list order by match_date DESC"
    cursor.execute(name_sql)
    results = cursor.fetchall()
    db.commit()
    db.close()
    for i in results:
        j = list(i)
        matchs.append(j)
    data['match_list'] = matchs
    response = HttpResponse(json.dumps(data))
    return response


def playercn_video(request):
    matchname = request.GET.get('matchname')
    videoname = request.GET.get('videoname')
    data = {}
    nl = videoname.split('-')
    if len(nl) == 3:
        name_a = nl[0]
        name_b = nl[2]
    elif len(nl) == 5:
        name_a = '-'.join([nl[0], nl[1]])
        name_b = '-'.join([nl[3], nl[4]])
    result = FormShipinReport.objects.filter(Q(query_saishimingcheng=matchname) & Q(single_line_shipinmingcheng__contains=videoname))[0]
    videoid = result.videoid
    if videoid == None:
        videoid = ''
    db = pymysql.connect('cba-db.hbang.com.cn', 'admin', 'Passw0rd235', 'cba')
    cursor = db.cursor()
    sql = "select * from finish where videoid='" + videoid + "'"
    cursor.execute(sql)
    result1 = cursor.fetchall()
    db.commit()
    db.close()
    db = pymysql.connect('cba-db.hbang.com.cn', 'admin', 'Passw0rd235', 'cba')
    cursor = db.cursor()
    sql_score = "select * from video_score where videoid='" + videoid + "'"
    cursor.execute(sql_score)
    result_score = cursor.fetchall()
    db.commit()
    db.close()
    a_h_d = 0
    a_w_d = 0
    a_f_d = 0
    a_q_d = 0
    a_w_s = 0
    a_z_s = 0
    a_b_s = 0
    b_h_d = 0
    b_w_d = 0
    b_f_d = 0
    b_q_d = 0
    b_w_s = 0
    b_z_s = 0
    b_b_s = 0

    for i in result_score:
        if i[8] == name_a:
            if i[9] == '后场直接得分':
                a_h_d += 1
            elif i[9] == '网前直接得分':
                a_w_d += 1
            elif i[9] == '发球直接得分':
                a_f_d += 1
            elif i[9] == '其他直接得分':
                a_q_d += 1
            elif i[9] == '对方网前失误':
                a_w_s += 1
            elif i[9] == '对方主动失误':
                a_z_s += 1
            elif i[9] == '对方被动失误':
                a_b_s += 1
        elif i[8] in name_b:
            if i[9] == '后场直接得分':
                b_h_d += 1
            elif i[9] == '网前直接得分':
                b_w_d += 1
            elif i[9] == '发球直接得分':
                b_f_d += 1
            elif i[9] == '其他直接得分':
                b_q_d += 1
            elif i[9] == '对方网前失误':
                b_w_s += 1
            elif i[9] == '对方主动失误':
                b_z_s += 1
            elif i[9] == '对方被动失误':
                b_b_s += 1
    wl_a = [name_a, a_h_d, a_w_d, a_f_d, a_q_d, a_w_s, a_z_s, a_b_s]
    wl_b = [name_b, b_h_d, b_w_d, b_f_d, b_q_d, b_w_s, b_z_s, b_b_s]

    data['videourl'] = result.url
    if len(result1) > 0:
        score_zip_url = result1[0][2]
        win_lost_zip_url = result1[0][3]
        start_zip_url = result1[0][4]
        end_zip_url = result1[0][5]
        data['score_zip_url'] = score_zip_url
        data['win_lost_zip_url'] = win_lost_zip_url
        data['start_zip_url'] = start_zip_url
        data['end_zip_url'] = end_zip_url

    else:
        data['score_zip_url'] = ''
        data['win_lost_zip_url'] = ''
        data['start_zip_url'] = ''
        data['end_zip_url'] = ''

    if len(result_score) > 0:
        data['wl_a'] = wl_a
        data['wl_b'] = wl_b
    else:
        data['wl_a'] = ''
        data['wl_b'] = ''
    response = HttpResponse(json.dumps(data))
    return response


def is_number(s):
    try:
        float(s)                    # 如果能转换float，说明是个数字
        return True
    except ValueError:
        pass                        # 占位符

    try:
        import unicodedata          # 引入Unicodedata模块
        unicodedata.numeric(s)      # 如果能转成numeric，说明是个数字
        return True
    except (TypeError,ValueError):
        pass

    return False


def upload(request):
    if request.method == 'POST':
        data = {}
        data["msg"] = '上传成功！'
        file_status = '处理完成'
        a = request.FILES.get('file')
        username = request.POST.get('username')
        fileclass = request.POST.get('fileclass')
        uploaddate = time.strftime("%Y-%m-%d", time.localtime())
        # route = 'D:\\9\\index\\' + a.name
        route = '/upload_files/' + a.name
        destination = open(route, 'wb+')
        for chunk in a.chunks():
            destination.write(chunk)
        destination.close()
        if fileclass == '训练量':
            if route.split('.')[-1] == 'xlsx' or route.split('.')[-1] == 'xls':
                xll_data = xlrd.open_workbook(route)
                table = xll_data.sheet_by_index(0)
                rowNum = table.nrows
                colNum = table.ncols
                l1 = []
                for i in range(0, rowNum):
                    l2 = []
                    for j in range(colNum):
                        value = table.cell_value(i, j)
                        l2.append(value)
                    l1.append(l2)
                t = l1[1][0].split('——')[0].replace(' ', '')
                t1 = re.findall(r'\d+', t)
                t = t1[0] + '-' + t1[1].zfill(2) + '-' + t1[2].zfill(2)
                dates = []
                tt = 0
                begin = t
                dt = datetime.datetime.strptime(begin, "%Y-%m-%d")
                date = begin[:]
                while tt < 7:
                    dates.append(date)
                    dt = dt + timedelta(1)
                    date = dt.strftime("%Y-%m-%d")
                    tt += 1
                date_all = []
                for date in dates:
                    t1 = date.split('-')
                    date = [str(int(t1[2])), '-', calendar.month_abbr[int(t1[1])], '-', t1[0]]
                    date = ''.join(date)
                    date_all.append(date)
                ll = []
                xm = ['技术（技战术）', '技术（小技术）', '体能（力量）', '体能（专项）']
                for i in l1[4:]:
                    for j in date_all:
                        m = date_all.index(j) * 4
                        for n in range(1, 5):
                            x = n + m
                            if isinstance(i[x], float):
                                d1 = ['', i[0].replace(' ', ''), j, xm[n - 1], int(i[x])]
                                ll.append(d1)
                            elif i[x] == '':
                                i[x] = '0'
                                d1 = ['', i[0].replace(' ', ''), j, xm[n - 1], int(i[x])]
                                ll.append(d1)
                for i in ll:
                    try:
                        body = {
                            'form_YunDongYuan': i[1],
                            'Number_XunLianShiChang': i[4],
                            'Date_field_XunLianRiQi': i[2],
                            'Dropdown_XunLianKeMu': i[3],
                        }
                        u1 = 'https://creator.zoho.com.cn/api/chunzhou.jia2/json/xunlian/form/form_XunLianLiang/record/add/'
                        params = {
                            'authtoken': 'd51ecfa14f98e8f14c91ac894bf8e7d4',
                            'scope': 'creatorapi',
                            'raw': "true",
                        }
                        req = requests.post(u1, params={**params, **body})
                        req_json = req.json()
                        # print(req_json)
                        if req_json['formname'][1]['operation'][1]['status'] != 'Success':
                            file_status = '内容有错误'
                    except Exception as e:
                        pass
            else:
                data["msg"] = '文件格式错误，请修改为xlsx或xls格式。'
        elif fileclass == '训练计划':
            if route.split('.')[-1] == 'docx':
                data_all = []
                docStr = Document(route)
                paragraph = docStr.paragraphs[2]
                parStr = paragraph.text
                l = parStr.split(' ')
                ll = []
                for i in l:
                    if i not in ll:
                        ll.append(i)
                y1 = re.findall(r'\d+', ll[0])
                t = y1[0] + '-' + y1[1].zfill(2) + '-' + y1[2].zfill(2)
                dates = []
                tt = 0
                begin = t
                dt = datetime.datetime.strptime(begin, "%Y-%m-%d")
                date = begin[:]
                while tt < 7:
                    dates.append(date)
                    dt = dt + timedelta(1)
                    date = dt.strftime("%Y-%m-%d")
                    tt += 1
                date_all = []
                for date in dates:
                    t1 = date.split('-')
                    date = [str(int(t1[2])), '-', calendar.month_abbr[int(t1[1])], '-', t1[0]]
                    date = ''.join(date)
                    date_all.append(date)
                zhubie = ll[4].split('：')[1][:2] + '一队'
                zhujiaolian = ll[6]
                jiaolianzu = [ll[8], ll[9], ll[10]]
                numTables = docStr.tables
                table = numTables[1]
                z1 = [table.cell(2, 1).text.split('\n'), table.cell(3, 1).text.split('\n'),
                      table.cell(4, 1).text.split('\n')]
                z2 = [table.cell(2, 2).text.split('\n'), table.cell(3, 2).text.split('\n'),
                      table.cell(4, 2).text.split('\n')]
                z3 = [table.cell(2, 3).text.split('\n'), table.cell(3, 3).text.split('\n'),
                      table.cell(4, 3).text.split('\n')]
                z4 = [table.cell(2, 4).text.split('\n'), table.cell(3, 4).text.split('\n'),
                      table.cell(4, 4).text.split('\n')]
                z5 = [table.cell(2, 5).text.split('\n'), table.cell(3, 5).text.split('\n'),
                      table.cell(4, 5).text.split('\n')]
                z6 = [table.cell(2, 6).text.split('\n'), table.cell(3, 6).text.split('\n'),
                      table.cell(4, 6).text.split('\n')]
                z7 = [table.cell(2, 7).text.split('\n'), table.cell(3, 7).text.split('\n'),
                      table.cell(4, 7).text.split('\n')]
                zz = [z1, z2, z3, z4, z5, z6, z7]
                for i in range(7):
                    xunlianriqi = date_all[i]
                    zjl = zhujiaolian
                    jlz = jiaolianzu
                    dd = zz[i]
                    for j in dd:
                        j = [x for x in j if x != '']
                        j = [x for x in j if x != ' ']
                        if len(j) > 2:
                            start_time = j[1]
                            start_time = start_time.replace('：', ':')
                            jj = j[2:]
                            for kc in jj:
                                tn = start_time.split(':')
                                hour = jj.index(kc)
                                s_time = str(int(tn[0]) + hour) + ':' + tn[1] + ':00'
                                e_time = str(int(tn[0]) + hour + 1) + ':' + tn[1] + ':00'
                                data_one = [xunlianriqi, s_time, e_time, zjl, jlz, zhubie, '未开始', kc]
                                data_all.append(data_one)
                for i in data_all[0:1]:
                    l1 = i
                    try:
                        body = {
                            'Date_field_XunLianRi': l1[0],
                            'Time_KaiShi': l1[1],
                            'Time_JieShu': l1[2],
                            'form_ZhiYuan_ZhuJiaoLian': l1[3],
                            'form_ZhiYuan_JiaoLianZu': l1[4],
                            'form_ZuZhiJiaGou': l1[5],
                            'Dropdown_BianGeng': l1[6],
                            'form_KeCheng': l1[7],
                        }

                        u1 = 'https://creator.zoho.com.cn/api/chunzhou.jia2/json/xunlian/form/form_XunLianJiHua/record/add/'
                        params = {
                            'authtoken': 'd51ecfa14f98e8f14c91ac894bf8e7d4',
                            'scope': 'creatorapi',
                            'raw': "true",
                        }
                        req = requests.post(u1, params={**params, **body})
                        req_json = req.json()
                    except Exception as e:
                        pass
            else:
                data["msg"] = '文件格式错误，请修改为docx格式。'
        elif fileclass == '周计划':
            if route.split('.')[-1] == 'docx':
                try:
                    docStr = Document(route)
                    paragraph = docStr.paragraphs[2]
                    parStr = paragraph.text
                    l = parStr.split(' ')
                    ll = []
                    for i in l:
                        if i not in ll:
                            ll.append(i)
                    y1 = re.findall(r'\d+', ll[0])
                    y2 = re.findall(r'\d+', ll[1])
                    start = y1[2] + '-' + calendar.month_abbr[int(y1[1])] + '-' + y1[0]
                    end = y2[2] + '-' + calendar.month_abbr[int(y2[1])] + '-' + y2[0]
                    zhubie = ll[4].split('：')[1][:2] + '一队'
                    zhujiaolian = ll[6]
                    jiaolianzu = [ll[8].split('：')[1], ll[9], ll[10]]
                    numTables = docStr.tables
                    zdsx = numTables[0].cell(0, 0).text.split('：')[1].replace('\n', '')
                    table = numTables[1]
                    ydls = [table.cell(6, 4).text, table.cell(7, 4).text, table.cell(8, 4).text]
                    qds = [table.cell(6, 5).text, table.cell(7, 5).text, table.cell(8, 5).text]
                    ydl = ''
                    qd = ''
                    for i in ydls:
                        i = i.replace(' ', '')
                        if i[-1] == '*':
                            ydl = i[:2]
                    for i in qds:
                        i = i.replace(' ', '')
                        if i[-1] == '*':
                            qd = i[:4]
                    yqmd = table.cell(6, 1).text.split('：')[1].replace('\n', '')
                    cishu = table.cell(6, 7).text.split('：')[-1]
                    l1 = [start, end, zhubie, zhujiaolian, jiaolianzu, zdsx, yqmd, cishu, ydl, qd]

                    body = {
                        'Date_field_KaiShi': l1[0],
                        'Date_field_JieShu': l1[1],
                        'form_ZuZhiJiaGou': l1[2],
                        'form_ZhiYuan_ZhuJiaoLian': l1[3],
                        'form_ZhiYuan_JiaoLianZu': l1[4],
                        'Single_Line_ZhiDaoSiXiang': l1[5],
                        'Single_Line_YaoQiuJiMuDi': l1[6],
                        'Number_CiShu': l1[7],
                        'Dropdown_YunDongLiang': l1[8],
                        'Dropdown_QiangDu': l1[9],
                    }
                    u1 = 'https://creator.zoho.com.cn/api/chunzhou.jia2/json/xunlian/form/form_ZhouJiHua/record/add/'
                    params = {
                        'authtoken': 'd51ecfa14f98e8f14c91ac894bf8e7d4',
                        'scope': 'creatorapi',
                        'raw': "true",
                    }
                    req = requests.post(u1, params={**params, **body})
                except Exception as e:
                    data["msg"] = '上传失败，内容错误。'
            else:

                data["msg"] = '文件格式错误，请修改为docx格式。'
        elif fileclass == '基础体能':
            if route.split('.')[-1] == 'xlsx' or route.split('.')[-1] == 'xls':
                try:
                    file = route
                    xlrd_data = xlrd.open_workbook(file)
                    table = xlrd_data.sheet_by_index(0)

                    rowNum = table.nrows
                    colNum = table.ncols

                    projects = ['form_BMI', 'form_ZuoWeiTiQianQu', 'form_30M', 'form_ChuiZhiZongTiao', 'form_ShenDun_MAX',
                                'form_WoTui_MAX',
                                'form_YinTiXiangShang', 'form_FuJiNaiLi', 'form_BeiJiNaiLi', 'form_3000M']
                    l1 = []
                    for i in range(0, rowNum):
                        l2 = []
                        for j in range(colNum):
                            value = table.cell_value(i, j)
                            l2.append(value)
                        l1.append(l2)
                    t = datetime.datetime.now()
                    date = ''.join([str(t.day), '-', calendar.month_abbr[int(t.month)], '-', str(t.year)])
                    n = 6
                    d1 = {
                        'form_BMI': [],
                        'form_ZuoWeiTiQianQu': [],
                        'form_30M': [],
                        'form_ChuiZhiZongTiao': [],
                        'form_ShenDun_MAX': [],
                        'form_WoTui_MAX': [],
                        'form_YinTiXiangShang': [],
                        'form_FuJiNaiLi': [],
                        'form_BeiJiNaiLi': [],
                        'form_3000M': []
                    }
                    while n < len(l1):
                        bmi_pinfen = l1[n + 1][3]
                        form_ZuoWeiTiQianQu = l1[n][4]
                        form_ZuoWeiTiQianQu_pinfen = l1[n + 1][4]
                        form_30M = l1[n][5]
                        form_30M_pinfen = l1[n + 1][5]
                        form_ChuiZhiZongTiao = l1[n][6]
                        form_ChuiZhiZongTiao_pinfen = l1[n + 1][6]
                        form_ShenDun_MAX = l1[n][7]
                        form_ShenDun_MAX_pinfen = l1[n + 1][7]
                        form_WoTui_MAX = l1[n][8]
                        form_WoTui_MAX_pinfen = l1[n + 1][8]
                        form_YinTiXiangShang = l1[n][9]
                        form_YinTiXiangShang_pinfen = l1[n + 1][9]
                        form_FuJiNaiLi = l1[n][10]
                        form_FuJiNaiLi_pinfen = l1[n + 1][10]
                        form_BeiJiNaiLi = l1[n][11]
                        form_BeiJiNaiLi_pinfen = l1[n + 1][11]
                        form_3000M = l1[n][12]
                        form_3000M_pinfen = l1[n + 1][12]
                        if is_number(bmi_pinfen) == False:
                            bmi_pinfen = '0'
                        if is_number(form_ZuoWeiTiQianQu) == False:
                            form_ZuoWeiTiQianQu = '0'
                        if is_number(form_ZuoWeiTiQianQu_pinfen) == False:
                            form_ZuoWeiTiQianQu_pinfen = '0'
                        if is_number(form_30M) == False:
                            form_30M = '0'
                        if is_number(form_30M_pinfen) == False:
                            form_30M_pinfen = '0'
                        if is_number(form_ChuiZhiZongTiao) == False:
                            form_ChuiZhiZongTiao = '0'
                        if is_number(form_ChuiZhiZongTiao_pinfen) == False:
                            form_ChuiZhiZongTiao_pinfen = '0'
                        if is_number(form_ShenDun_MAX) == False:
                            form_ShenDun_MAX = '0'
                        if is_number(form_ShenDun_MAX_pinfen) == False:
                            form_ShenDun_MAX_pinfen = '0'
                        if is_number(form_WoTui_MAX) == False:
                            form_WoTui_MAX = '0'
                        if is_number(form_WoTui_MAX_pinfen) == False:
                            form_WoTui_MAX_pinfen = '0'
                        if is_number(form_YinTiXiangShang) == False:
                            form_YinTiXiangShang = '0'
                        if is_number(form_YinTiXiangShang_pinfen) == False:
                            form_YinTiXiangShang_pinfen = '0'
                        if is_number(form_FuJiNaiLi) == False:
                            form_FuJiNaiLi = '0'
                        if is_number(form_FuJiNaiLi_pinfen) == False:
                            form_FuJiNaiLi_pinfen = '0'
                        if is_number(form_BeiJiNaiLi) == False:
                            form_BeiJiNaiLi = '0'
                        if is_number(form_BeiJiNaiLi_pinfen) == False:
                            form_BeiJiNaiLi_pinfen = '0'
                        if is_number(form_3000M) == False:
                            form_3000M = '0'
                        if is_number(form_3000M_pinfen) == False:
                            form_3000M_pinfen = '0'
                        bmi_pinfen = int(bmi_pinfen)
                        form_ZuoWeiTiQianQu = form_ZuoWeiTiQianQu
                        form_ZuoWeiTiQianQu_pinfen = int(form_ZuoWeiTiQianQu_pinfen)
                        form_30M = form_30M
                        form_30M_pinfen = int(form_30M_pinfen)
                        form_ChuiZhiZongTiao = form_ChuiZhiZongTiao
                        form_ChuiZhiZongTiao_pinfen = int(form_ChuiZhiZongTiao_pinfen)
                        form_ShenDun_MAX = int(form_ShenDun_MAX)
                        form_ShenDun_MAX_pinfen = int(form_ShenDun_MAX_pinfen)
                        form_WoTui_MAX = form_WoTui_MAX
                        form_WoTui_MAX_pinfen = int(form_WoTui_MAX_pinfen)
                        form_YinTiXiangShang = int(form_YinTiXiangShang)
                        form_YinTiXiangShang_pinfen = int(form_YinTiXiangShang_pinfen)
                        form_FuJiNaiLi = form_FuJiNaiLi
                        form_FuJiNaiLi_pinfen = int(form_FuJiNaiLi_pinfen)
                        form_BeiJiNaiLi = form_BeiJiNaiLi
                        form_BeiJiNaiLi_pinfen = int(form_BeiJiNaiLi_pinfen)
                        form_3000M = form_3000M
                        form_3000M_pinfen = int(form_3000M_pinfen)
                        d_bmi = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShiRiQi': date, 'Number_PingFen': bmi_pinfen}
                        d_ZuoWeiTiQianQu = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                            'Decimal_CeShiJieGuo': form_ZuoWeiTiQianQu,
                                            'Dropdown_PingFen': form_ZuoWeiTiQianQu_pinfen}
                        d_30M = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date, 'Decimal_CeShiJieGuo': form_30M,
                                 'Dropdown_PingFen': form_30M_pinfen}
                        d_ChuiZhiZongTiao = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                             'Decimal_CeShiJieGuo': form_ChuiZhiZongTiao,
                                             'Dropdown_PingFen': form_ChuiZhiZongTiao_pinfen}
                        d_ShenDun_MAX = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                         'Number_CeShiJieGuo': form_ShenDun_MAX,
                                         'Dropdown_PingFen': form_ShenDun_MAX_pinfen}
                        d_WoTui_MAX = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                       'Decimal_CeShiJieGuo': form_WoTui_MAX, 'Dropdown_PingFen': form_WoTui_MAX_pinfen}
                        d_YinTiXiangShang = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                             'Number_CeShiJieGuo': form_YinTiXiangShang,
                                             'Dropdown_PingFen': form_YinTiXiangShang_pinfen}
                        d_FuJiNaiLi = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                       'Decimal_CeShiJieGuo': form_FuJiNaiLi, 'Dropdown_PingFen': form_FuJiNaiLi_pinfen}
                        d_BeiJiNaiLi = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                        'Decimal_CeShiJieGuo': form_BeiJiNaiLi, 'Dropdown_PingFen': form_BeiJiNaiLi_pinfen}
                        d_3000M = {'form_YunDongYuan': l1[n][0], 'Date_field_CeShi': date,
                                   'Decimal_CeShiJieGuo': form_3000M, 'Dropdown_PingFen': form_3000M_pinfen}
                        d1['form_BMI'].append(d_bmi)
                        d1['form_ZuoWeiTiQianQu'].append(d_ZuoWeiTiQianQu)
                        d1['form_30M'].append(d_30M)
                        d1['form_ChuiZhiZongTiao'].append(d_ChuiZhiZongTiao)
                        d1['form_ShenDun_MAX'].append(d_ShenDun_MAX)
                        d1['form_WoTui_MAX'].append(d_WoTui_MAX)
                        d1['form_YinTiXiangShang'].append(d_YinTiXiangShang)
                        d1['form_FuJiNaiLi'].append(d_FuJiNaiLi)
                        d1['form_BeiJiNaiLi'].append(d_BeiJiNaiLi)
                        d1['form_3000M'].append(d_3000M)
                        n += 2
                    for i in projects:
                        vals = d1[i]
                        for j in vals:
                            try:
                                body = j
                                u1 = 'https://creator.zoho.com.cn/api/chunzhou.jia2/json/tineng/form/' + i + '/record/add/'
                                params = {
                                    'authtoken': 'd51ecfa14f98e8f14c91ac894bf8e7d4',
                                    'scope': 'creatorapi',
                                    'raw': "true",
                                }
                                req = requests.post(u1, params={**params, **body})
                                req_json = req.json()
                            except Exception as e:
                                pass
                except Exception as e:
                    data['msg'] = '上传失败，数据错误！'

            else:
                data["msg"] = '文件格式错误，请修改为xlsx或xls格式。'
        elif fileclass == '专项体能':
            pass
        elif fileclass == '生理生化':
            if route.split('.')[-1] == 'xlsx' or route.split('.')[-1] == 'xls':
                try:
                    data_slsh = xlrd.open_workbook(route)
                    table = data_slsh.sheet_by_index(0)
                    rowNum = table.nrows
                    colNum = table.ncols
                    l1 = []
                    for i in range(0, rowNum):
                        l2 = []
                        for j in range(colNum):
                            value = table.cell_value(i, j)
                            l2.append(value)
                        l1.append(l2)
                    n = 1
                    data_all = []
                    while n < len(l1):
                        t = l1[n][0].split('.')
                        y = t[0]
                        m = t[1]
                        d = t[2]
                        date = ''.join([d, '-', calendar.month_abbr[int(m)], '-', y])
                        hxb = l1[n][3]
                        bxb = l1[n][4]
                        xqyj = l1[n][5]
                        xhdb = l1[n][6]
                        jsjm = l1[n][7]
                        gt = l1[n][8]
                        xns = l1[n][9]
                        pzc = l1[n][10]
                        if is_number(hxb) == False:
                            hxb = '0'
                        if is_number(bxb) == False:
                            bxb = '0'
                        if is_number(xqyj) == False:
                            xqyj = '0'
                        if is_number(xhdb) == False:
                            xhdb = '0'
                        if is_number(jsjm) == False:
                            jsjm = '0'
                        if is_number(gt) == False:
                            gt = '0'
                        if is_number(xns) == False:
                            xns = '0'
                        if is_number(pzc) == False:
                            pzc = '0'

                        d_hxb = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '红细胞',
                                 'Decimal_CeShiJieGuo': hxb}
                        d_bxb = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '白细胞',
                                 'Decimal_CeShiJieGuo': bxb}
                        d_xqyj = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '血球压积',
                                  'Decimal_CeShiJieGuo': xqyj}
                        d_xhdb = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '血红蛋白',
                                  'Decimal_CeShiJieGuo': xhdb}
                        d_jsjm = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '肌酸激酶',
                                  'Decimal_CeShiJieGuo': jsjm}
                        d_gt = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '睾酮',
                                'Decimal_CeShiJieGuo': gt}
                        d_xns = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '血尿素',
                                 'Decimal_CeShiJieGuo': xns}
                        d_pzc = {'form_YunDongYuan': l1[n][1], 'Date_field_CeShiRiQi': date, 'Dropdown_CeShiXiangMu': '皮质醇',
                                 'Decimal_CeShiJieGuo': pzc}
                        data_all.append(d_hxb)
                        data_all.append(d_bxb)
                        data_all.append(d_xqyj)
                        data_all.append(d_xhdb)
                        data_all.append(d_jsjm)
                        data_all.append(d_gt)
                        data_all.append(d_xns)
                        data_all.append(d_pzc)
                        n += 1

                    for j in data_all:
                        try:
                            body = j
                            u1 = 'https://creator.zoho.com.cn/api/chunzhou.jia2/json/myapplication/form/form_ShengLiShengHua/record/add/'
                            params = {
                                'authtoken': 'd51ecfa14f98e8f14c91ac894bf8e7d4',
                                'scope': 'creatorapi',
                                'raw': "true",
                            }
                            req = requests.post(u1, params={**params, **body})
                            req_json = req.json()
                        except Exception as e:
                            pass
                except Exception as e:
                    data['msg'] = '上传失败，数据错误！'
            else:
                data["msg"] = '文件格式错误，请修改为xlsx或xls格式。'
        elif fileclass == 'FMS':
            pass
        aws_upload = 'aws s3 cp ' + route + ' s3://video.hbang.com.cn/upload_files/' + a.name
        os.system(aws_upload)
        if data['msg'] == '上传成功！':
            up_data = file_list()
            up_data.file_name = a.name
            up_data.file_class = fileclass
            up_data.upload_user = username
            up_data.upload_DATE = uploaddate
            up_data.status = file_status
            up_data.save()

        response = HttpResponse(json.dumps(data))
        return response
    if request.method == 'GET':
        data = {'files': []}
        start = datetime.date.today() - relativedelta(months=3)
        results = file_list.objects.filter(upload_DATE__gte=start)
        for i in results:
            date = str(i.upload_DATE.year) + '-' + str(i.upload_DATE.month) + '-' + str(i.upload_DATE.day)
            dd = [i.file_name, i.file_class, i.upload_user, date, i.status]
            data['files'].append(dd)
        data['files'].reverse()
        response = HttpResponse(json.dumps(data))
        return response


def update(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        new_pwd = request.POST.get('pass')
        dd = video_user.objects.get(username=username)
        dd.password = new_pwd
        dd.save()
        data = {'data': '修改成功！'}
        response = HttpResponse(json.dumps(data))
        return response

